from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_UNDERLINE
from docx.enum.section import WD_SECTION
from docx.shared import Cm

def create_numbered_tables(start_num=1, end_num=500):
    doc = Document()
    
    # Установка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Параметры таблицы
    rows_per_page = 12
    cols = 4
    numbers_per_page = rows_per_page * cols
    
    # Вычисляем количество необходимых страниц
    current_num = start_num
    
    while current_num <= end_num:
        # Создаем таблицу для текущей страницы
        table = doc.add_table(rows=rows_per_page, cols=cols)
        table.style = 'Table Grid'
        
        # Устанавливаем ширину столбцов
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(1.5)
        
        # Заполняем таблицу числами
        for row in range(rows_per_page):
            for col in range(cols):
                if current_num <= end_num:
                    cell = table.cell(row, col)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(str(current_num))
                    
                    # Форматирование
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.underline = WD_UNDERLINE.SINGLE
                    
                    # Центрирование текста
                    paragraph.alignment = 1  # 1 = CENTER
                    
                    current_num += 1
        
        # Добавляем разрыв страницы, если это не последняя страница
        if current_num <= end_num:
            doc.add_page_break()
    
    # Сохраняем документ
    doc.save('numbered_pages.docx')

# Создаем документ с числами от 1 до 500
create_numbered_tables(1, 500)