from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches

def create_nested_tables(start_num=1, end_num=500):
    # Create document
    doc = Document()
    
    # Parameters for inner tables
    inner_cols = 4
    inner_rows = 12
    numbers_per_cell = inner_cols * inner_rows
    
    # Calculate how many outer cells we need
    total_numbers = end_num - start_num + 1
    outer_cells_needed = total_numbers // numbers_per_cell
    if total_numbers % numbers_per_cell:
        outer_cells_needed += 1
    
    # Calculate outer table dimensions
    outer_cols = 4
    outer_rows = (outer_cells_needed + outer_cols - 1) // outer_cols
    
    # Create outer table
    outer_table = doc.add_table(rows=outer_rows, cols=outer_cols)
    outer_table.style = 'Table Grid'
    
    current_number = start_num
    
    # Fill each cell of outer table with inner table
    for outer_row in range(outer_rows):
        for outer_col in range(outer_cols):
            if current_number <= end_num:
                outer_cell = outer_table.cell(outer_row, outer_col)
                
                # Create inner table in this cell
                inner_table = outer_cell.add_table(rows=inner_rows, cols=inner_cols)
                inner_table.style = 'Table Grid'
                
                # Fill inner table with numbers
                for inner_row in range(inner_rows):
                    for inner_col in range(inner_cols):
                        if current_number <= end_num:
                            inner_cell = inner_table.cell(inner_row, inner_col)
                            paragraph = inner_cell.paragraphs[0]
                            run = paragraph.add_run(str(current_number))
                            
                            # Format number
                            run.font.size = Pt(10)
                            run.font.underline = WD_UNDERLINE.SINGLE
                            run.font.bold = True
                            
                            current_number += 1
    
    # Save document
    doc.save('nested_tables.docx')

# Create document with nested tables
create_nested_tables()